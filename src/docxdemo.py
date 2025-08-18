from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn

def calculate_scaled_dimensions(original_width, original_height, available_width, available_height, max_ratio=0.7):
    """
    计算保持宽高比的缩放尺寸
    :param max_ratio: 图片占可用区域的最大比例（0.7表示70%）
    """
    # 计算可用区域的70%空间
    target_width = available_width * max_ratio
    target_height = available_height * max_ratio
    
    # 计算宽高缩放比例
    width_ratio = target_width / original_width
    height_ratio = target_height / original_height
    
    # 选择较小的缩放比例以确保图片完整显示
    scale_ratio = min(width_ratio, height_ratio)
    
    return int(original_width * scale_ratio), int(original_height * scale_ratio)

def create_word_demo(file_path):
    # 创建文档并设置窄边距（单位：厘米）
    doc = Document()
    section = doc.sections[0]
    
    # 设置页边距（可修改为需要的最小值）
    left_margin = Cm(1.5)
    right_margin = Cm(1.5)
    top_margin = Cm(1.2)
    bottom_margin = Cm(1.2)
    
    section.left_margin = left_margin
    section.right_margin = right_margin
    section.top_margin = top_margin
    section.bottom_margin = bottom_margin

    # 计算页面可用宽度（A4纸宽度21cm - 左右边距）
    page_width = Cm(21) - left_margin - right_margin
    page_height = Cm(29.7) - top_margin - bottom_margin  # A4纸高度29.7cm

    # 添加标题
    title = doc.add_heading('智能图片缩放示例', level=1)
    title_run = title.runs[0]
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title_run.font.size = Pt(20)

    # 添加说明文字
    doc.add_paragraph("以下图片将自动缩放至页面70%区域内显示:", style='Body Text')
    
    # 添加3个空行
    doc.add_paragraph("\n" * 3)

    # 添加图片（自动缩放）
    try:
        img_path = 'example.png'  # 替换为实际图片路径
        
        # 获取图片原始尺寸（单位：EMU）
        from docx.image.image import Image
        img = Image.from_file(img_path)
        original_width, original_height = img.width, img.height
        
        # 计算缩放后尺寸（转换为英寸单位）
        scaled_width, scaled_height = calculate_scaled_dimensions(
            original_width, 
            original_height,
            page_width,
            page_height
        )
        
        # 添加图片到文档（使用计算后的尺寸）
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        inline_shape = run.add_picture(img_path, width=Inches(scaled_width/914400), height=Inches(scaled_height/914400))
        
        # 添加图片信息说明
        info = doc.add_paragraph()
        info.add_run(f"原始尺寸: {original_width/914400:.1f}×{original_height/914400:.1f}英寸 | ")
        info.add_run(f"缩放后: {scaled_width/914400:.1f}×{scaled_height/914400:.1f}英寸 | ")
        info.add_run(f"占页面区域: 70%").bold = True
        info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    except Exception as e:
        doc.add_paragraph(f'[图片加载失败: {str(e)}]', style='Body Text')

    # 保存文档
    doc.save(file_path)
    print(f"文档已保存到: {file_path}")
    print(f"页面边距: 左右{left_margin.cm}cm 上下{top_margin.cm}cm")
    print(f"图片自动缩放至页面70%区域内")

# 使用示例
file_path = r'./advanced_demo.docx'
create_word_demo(file_path)