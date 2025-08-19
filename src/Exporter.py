from docx import Document
from docx.image.image import Image
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.shared import qn

import sys
import os
from datetime import datetime

import DataManagement

def date():
    _now = datetime.now()
    DATE = _now.strftime("%Y-%m-%d %H-%M-%S")
    return DATE


class PageMonitor():
    def __init__(self, doc):
        self.lines_per_page = None
        self.remains = None
        self.initial(doc)

    def initial(self, doc):
        height = doc.PAGE_HEIGHT.inches * 72
        lines = int(height / doc.SIZE)
        self.lines_per_page = lines
        self.remains = lines

    def remain_lines(self):
        return self.remain_lines

    def reset(self):
        self.remains = self.lines_per_page
    
    def if_available(self, lines):
        return self.remains > int(lines * 1.2)
    
    def add_lines(self, lines):
        self.remains -= lines

class DOCX():
    def __init__(self):
        # constant
        # Emu
        self.LEFT_MARGIN = None
        self.RIGHT_MARGIN = None
        self.TOP_MARGIN = None
        self.BOTTOM_MARGIN = None
        self.PAGE_WIDTH = None
        self.PAGE_HEIGHT = None
        self.FONT = '宋体-简' if sys.platform == 'darwin' else '宋体'
        self.SIZE = 9

        self.doc = self.generate_docx()

        self.Monitor = PageMonitor(self)

    def generate_docx(self):
        # generate document
        doc = Document()
        section = doc.sections[0]

        # adjust margin
        self.LEFT_MARGIN = Cm(1.27)
        self.RIGHT_MARGIN = Cm(1.27)
        self.TOP_MARGIN = Cm(1.27)
        self.BOTTOM_MARGIN = Cm(1.27)

        section.left_margin = Emu(self.LEFT_MARGIN)
        section.right_margin = Emu(self.RIGHT_MARGIN)
        section.top_margin = Emu(self.TOP_MARGIN)
        section.bottom_margin = Emu(self.BOTTOM_MARGIN)

        # get available area for A4
        page_width = Cm(21) - self.LEFT_MARGIN - self.RIGHT_MARGIN
        page_height = Cm(29.7) - self.TOP_MARGIN - self.BOTTOM_MARGIN
        self.PAGE_WIDTH = Emu(page_width)
        self.PAGE_HEIGHT = Emu(page_height)

        # sigle line spacing
        style = doc.styles['Normal']
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(0)  # 段前间距0
        style.paragraph_format.space_after = Pt(0)   # 段后间距0

        return doc
    
    def new_page(self):
        self.doc.add_page_break()
    
    def add_blank_row(self):
        doc = self.doc
        paragraph = doc.add_paragraph(' ', style='Body Text')
        run = paragraph.runs[0]
        run.font.name = self.FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT)
        run.font.size = Pt(self.SIZE)

    def insert_paragraph(self, content, font, size):
        """
        insert a pure text paragraph
        """
        doc = self.doc
        paragraph = doc.add_paragraph(content, style='Body Text')
        run = paragraph.runs[0]
        run.font.name = font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        run.font.size = Pt(size)

        self.Monitor.add_lines(1)
        return True

    def calculate_scaled_dimensions(self, original_width, original_height, available_width, available_height, max_ratio):
        """
        calculate scaling ratio. Scaled image should occupy availabel area at max_ratio.
        input : EMU
        return : scaled_width, scaled_height (both EMU)
        """
        target_width = available_width * max_ratio
        target_height = available_height * max_ratio
        
        width_ratio = target_width / original_width
        height_ratio = target_height / original_height
        
        scale_ratio = min(width_ratio, height_ratio)
        
        return Emu(original_width * scale_ratio), Emu(original_height * scale_ratio)

    def insert_image(self, img_path, width, height):
        """
        insert an image to a single paragraph
        """
        try:
            if os.path.isfile(img_path):
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run()
                inline_image = run.add_picture(img_path, width=width, height=height)
            return True
        except Exception:
            return False
    
    def write_one_question(self, data, blank_row, ratio):
        # data: (ID, subject, source, page, number, times, keypoints, note, image_path)
        # question style:
        # for any line, only if info exists, write it
        # ID: ID, 科目:subject, 来源: source, 页码: page, 编号: number, 错误次数: times
        # 知识点: keypoints list
        # 备注: note
        (ID, subject, source, page, number, times, keypoints, note, image_path) = data
        
        # prepare text content
        line1 = "ID: {}, 科目: {}, 来源: {}, 页码: {}, 编号: {}, 错误次数: {}".format(ID, subject, source, page, number, times)
        lines_count = 1
        line2 = ''
        if keypoints:
            line2 = '知识点: ' + ' ; '.join(keypoints)
            lines_count += 1
        line3 = ''
        if note:
            line3 = '备注' + note
            lines_count += 1

        # prepare image content
        if not os.path.isfile(image_path):
            return False
        image = Image.from_file(image_path)
        image_width = image.width
        image_height = image.height
        width, height = self.calculate_scaled_dimensions(
            image_width, image_height, 
            self.PAGE_WIDTH, self.PAGE_HEIGHT,
            ratio
            )
        width = Inches(width.inches)
        height = Inches(height.inches)
        lines_count += int((height.inches * 72) / self.SIZE) + 3

        lines_count += blank_row

        # if remaining space is enough
        if not self.Monitor.if_available(lines_count):
            self.new_page()
            self.Monitor.reset()
        
        # update remaining space
        self.Monitor.add_lines(lines_count)
        
        # write text content
        lines = [line1, line2, line3]
        temp = []
        for item in lines:
            if item != '' and item != ' ':
                temp.append(item)
        for i in range(len(temp)):
            self.insert_paragraph(temp[i], self.FONT, self.SIZE)
        # insert image
        self.insert_image(image_path, width=width, height=height)
        # blank rows
        for i in range(blank_row):
            self.add_blank_row()

        return True
    
    def output(self, IDs, output_path, blank_row, ratio):
        disk = DataManagement.DiskController()
        # info line
        self.insert_paragraph(f'生成于{date()}', font=self.FONT, size=self.SIZE)
        self.add_blank_row()
        # questions
        for ID in IDs:
            ID, subject, source, times, image_path, keypoints, note, answer, page, number = disk.read_questions(ID)
            data = (ID, subject, source, page, number, times, keypoints, note, image_path)
            #print(data)
            if not self.write_one_question(data, blank_row, ratio):
                return False, ''
        # save
        if os.path.isdir(output_path):
            path = os.path.join(output_path, f'错题导出-{date()}.docx')
            self.doc.save(path)

        return True, path

if __name__ == '__main__':
    pass